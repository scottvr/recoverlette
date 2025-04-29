import argparse
import asyncio
import time
import os
import requests # Still potentially useful for fallback
import sys
import uuid
from pathlib import Path
import re
import io
import logging # Added for logging
import urllib.parse
from dataclasses import dataclass

# --- Load .env file ---
from dotenv import load_dotenv
load_dotenv()
# --- End Load .env file ---

# --- Add python-docx dependency ---
try:
    import docx
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    from docx.shared import RGBColor
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: The 'python-docx' library is required.", file=sys.stderr)
    print("Please install it using: pip install python-docx", file=sys.stderr)
    sys.exit(1)
# --- End python-docx dependency ---

# Authentication & SDK Core
from azure.identity import InteractiveBrowserCredential, TokenCachePersistenceOptions
from msgraph import GraphServiceClient
from msgraph.generated.models.o_data_errors.o_data_error import ODataError
from msgraph.generated.models.drive_item import DriveItem
from msgraph.generated.models.item_reference import ItemReference
from kiota_abstractions.base_request_configuration import RequestConfiguration
from kiota_abstractions.default_query_parameters import QueryParameters

# --- Configuration Loading ---
CLIENT_ID = os.getenv("RECOVERLETTE_CLIENT_ID")
TENANT_ID = os.getenv("RECOVERLETTE_TENANT_ID", "consumers")

if not CLIENT_ID:
    logging.critical("Error: Configuration variable RECOVERLETTE_CLIENT_ID is not set.")
    sys.exit(1)

SCOPES = ['Files.ReadWrite', 'User.Read']

# --- Configure Logging ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
az_identity_logger = logging.getLogger('azure.identity')
az_identity_logger.setLevel(logging.WARNING)

# --- End Configure Logging ---


# --- Helper Class for Argument Parsing ---
class DefineAction(argparse.Action):
    """Custom action to parse KEY=VALUE pairs for definitions."""
    def __call__(self, parser, namespace, values, option_string=None):
        definitions = getattr(namespace, self.dest, {}) or {}
        for value_pair in values:
            if '=' not in value_pair:
                raise argparse.ArgumentError(self, f"Invalid definition format: '{value_pair}'. Use KEY=VALUE.")
            key, value = value_pair.split('=', 1)
            definitions[key.strip()] = value.strip()
        setattr(namespace, self.dest, definitions)

# --- Authentication ---
auth_credential = None

async def get_authenticated_client() -> GraphServiceClient | None:
    """Creates and returns an authenticated GraphServiceClient."""
    global auth_credential
    # Use root logger for general script info
    logger = logging.getLogger(__name__)

    logger.info(f"Using Client ID: ***{CLIENT_ID[-4:]}")
    logger.info(f"Using Tenant ID: {TENANT_ID}")

    cache_options = TokenCachePersistenceOptions(name="recoverlette_cache")
    logger.debug("TokenCachePersistenceOptions created (name='recoverlette_cache').")

    credential = None
    try:
         logger.debug("Attempting to create InteractiveBrowserCredential...")
         credential = InteractiveBrowserCredential(
             client_id=CLIENT_ID,
             tenant_id=TENANT_ID,
             cache_persistence_options=cache_options
             )
         auth_credential = credential
         logger.info("InteractiveBrowserCredential created, persistent token cache enabled.")
    except Exception as e:
         logger.exception(f"Error creating credential object: {e}")
         return None

    logger.debug("Creating GraphServiceClient...")
    if not credential:
        logger.error("Credential object is None, cannot create GraphServiceClient.")
        return None
    client = GraphServiceClient(credentials=credential, scopes=SCOPES)
    logger.debug("GraphServiceClient created.")

    logger.info("Attempting authentication check (will use cache or trigger browser flow)...")
    try:
        request_config = RequestConfiguration(
            query_parameters = {'select': ['displayName']}
        )
        logger.debug("Attempting client.me.get() to verify authentication...")
        me_user = await client.me.get(request_configuration=request_config)
        logger.debug("client.me.get() call completed.")

        if me_user and me_user.display_name:
             logger.info(f"Authentication successful for user: {me_user.display_name}")
             return client
        elif me_user:
             logger.warning("Authentication check successful, but couldn't retrieve user display name.")
             return client
        else:
             logger.error("Authentication check call did not return a user object.")
             return None

    except ODataError as o_data_error:
        logger.error(f"Authentication or initial Graph call failed:")
        if o_data_error.error:
            logger.error(f"  Code: {o_data_error.error.code}")
            logger.error(f"  Message: {o_data_error.error.message}")
        logger.debug("ODataError details:", exc_info=True)
        return None
    except Exception as e:
        logger.exception(f"An unexpected error occurred during authentication or Graph call: {e}")
        response_body = getattr(e, 'response', None)
        if response_body is not None:
             try:
                  status = getattr(response_body, 'status_code', 'N/A')
                  body_text = getattr(response_body, 'text', '{}')
                  logger.error(f"  Underlying HTTP Status: {status}")
                  logger.error(f"  Underlying Response Body: {body_text}")
             except Exception as inner_e:
                  logger.error(f"  Could not extract full details from exception response object: {inner_e}")
        return None

# --- Placeholder Discovery ---
def find_placeholders_in_docx(content_bytes: bytes) -> set[str]:
    """Uses python-docx to find all unique placeholder keys [[KEY]]."""
    logger = logging.getLogger(__name__)
    found_keys = set()
    # *** CHANGED REGEX to use @...@ delimiter ***
    placeholder_pattern = re.compile(r"\[\[(.*?)\]\]") 
    
    logger.debug("Starting DOCX placeholder scan for [[...]].")
    try:
        doc_stream = io.BytesIO(content_bytes)
        document = docx.Document(doc_stream)
        # Iterate through paragraphs and tables (same structure as before)
        for para in document.paragraphs:
            for run in para.runs:
                 matches = placeholder_pattern.findall(run.text)
                 if matches:
                     logger.debug(f"  Found in paragraph run: {matches}")
                     for match in matches:
                         found_keys.add(match.strip()) # Store key without @ symbols
            # Check full paragraph text as fallback
            matches = placeholder_pattern.findall(para.text)
            if matches:
                logger.debug(f"  Found in paragraph text: {matches}")
                for match in matches:
                    found_keys.add(match.strip()) # Store key without @ symbols
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                         for run in para.runs:
                              matches = placeholder_pattern.findall(run.text)
                              if matches:
                                   logger.debug(f"  Found in table cell run: {matches}")
                                   for match in matches:
                                       found_keys.add(match.strip())
                         matches = placeholder_pattern.findall(para.text)
                         if matches:
                             logger.debug(f"  Found in table cell paragraph text: {matches}")
                             for match in matches:
                                 found_keys.add(match.strip())
    except Exception as e:
        logger.warning(f"Error parsing DOCX template to find placeholders: {e}", exc_info=True)
        logger.warning("Placeholder reporting might be incomplete.")
        return set()
    logger.debug(f"Placeholder scan finished. Found unique keys: {found_keys}")
    return found_keys



# --- Graph Operations ---
async def get_drive_item_details(client: GraphServiceClient, file_path: str) -> tuple[str | None, str | None, str | None]:
    """Gets the OneDrive item ID, parent folder ID, and drive ID."""
    # (Same as previous version)
    logger = logging.getLogger(__name__)
    item_id = None
    parent_folder_id = None
    drive_id = None
    logger.debug(f"Attempting to get drive item details for path: {file_path}")
    try:
        logger.debug("Getting user's default drive ID...")
        drive_info_config = RequestConfiguration(query_parameters={'select': ['id']})
        drive_info = await client.me.drive.get(request_configuration=drive_info_config)
        if not drive_info or not drive_info.id:
            logger.error("Could not retrieve user's default drive ID.")
            return None, None, None
        drive_id = drive_info.id
        logger.debug(f"Using Drive ID: {drive_id}")
        encoded_file_path = file_path.lstrip('/')
        encoded_file_path = encoded_file_path.replace("#", "%23").replace("?", "%3F")
        path_based_id = f"root:/{encoded_file_path}"
        logger.debug(f"Attempting to get item using path-based ID: '{path_based_id}' within drive {drive_id}")
        item_request_config = RequestConfiguration(
            query_parameters = {'select': ["id", "parentReference"]}
        )
        drive_item = await client.drives.by_drive_id(drive_id).items.by_drive_item_id(path_based_id).get(
             request_configuration=item_request_config
        )
        if drive_item and drive_item.id:
            item_id = drive_item.id
            if drive_item.parent_reference and drive_item.parent_reference.id:
                 parent_folder_id = drive_item.parent_reference.id
            else:
                 logger.debug(f"Item {item_id} has no parentReference, attempting to get root ID for drive {drive_id}.")
                 root_request_config = RequestConfiguration(query_parameters={'select': ["id"]})
                 root_item = await client.drives.by_drive_id(drive_id).root.get(request_configuration=root_request_config)
                 if root_item and root_item.id:
                     if item_id == root_item.id:
                          parent_folder_id = root_item.id
                          logger.debug(f"Item {item_id} appears to be the root folder.")
                     else:
                          logger.warning(f"Item {item_id} has no parentReference but is not the root folder.")
                          parent_folder_id = root_item.id
            if parent_folder_id:
                 logger.info(f"Found Item ID: {item_id}, Parent Folder ID: {parent_folder_id}, Drive ID: {drive_id}")
            else:
                 logger.warning(f"Found Item ID: {item_id}, Drive ID: {drive_id}, but could not determine Parent Folder ID.")
        else:
            logger.error(f"Could not retrieve item details using path-based ID '{path_based_id}'")
    except ODataError as o_data_error:
        logger.error(f"ODataError getting item details for {file_path}:")
        if o_data_error.error:
            logger.error(f"  Code: {o_data_error.error.code}")
            logger.error(f"  Message: {o_data_error.error.message}")
            response_status = getattr(o_data_error, 'response_status_code', None)
            if response_status == 404 or "itemNotFound" in (o_data_error.error.code or ""):
                 logger.error(f"  Hint: Item not found. Check if the path '{file_path}' is correct and exists in OneDrive. The path ID used was '{path_based_id}'.")
        logger.debug("ODataError details:", exc_info=True)
    except Exception as e:
        logger.exception(f"An unexpected error occurred getting item details for {file_path}: {e}")
    return item_id, parent_folder_id, drive_id


async def get_file_content(client: GraphServiceClient, drive_id: str, item_id: str) -> bytes | None:
    """Downloads file content for a given drive ID and item ID."""
    # (Same as previous version)
    logger = logging.getLogger(__name__)
    logger.debug(f"Attempting to download content for item ID: {item_id} in drive {drive_id}")
    try:
        content_result = await client.drives.by_drive_id(drive_id).items.by_drive_item_id(item_id).content.get()
        if isinstance(content_result, bytes):
            content_bytes = content_result
            if not content_bytes:
                 logger.warning(f"Downloaded content was empty for item {item_id}.")
            logger.info(f"Successfully downloaded content for item {item_id} ({len(content_bytes)} bytes).")
            return content_bytes
        elif hasattr(content_result, 'iter_bytes'):
             logger.debug("Content received as a stream, iterating...")
             content_bytes = b""
             async for chunk in content_result.iter_bytes():
                  content_bytes += chunk
             if not content_bytes:
                  logger.warning(f"Downloaded content stream was empty for item {item_id}.")
             logger.info(f"Successfully downloaded content stream for item {item_id} ({len(content_bytes)} bytes).")
             return content_bytes
        else:
             logger.error(f"Received unexpected type {type(content_result)} when downloading content for item {item_id}.")
             return None
    except ODataError as o_data_error:
        logger.error(f"Error downloading content for item {item_id}:")
        if o_data_error.error:
            logger.error(f"  Code: {o_data_error.error.code}")
            logger.error(f"  Message: {o_data_error.error.message}")
        logger.debug("ODataError details:", exc_info=True)
        return None
    except Exception as e:
        logger.exception(f"An unexpected error occurred downloading content for {item_id}: {e}")
        return None

# --- Placeholder Replacement ---
def replace_placeholders(content_bytes: bytes, defined_replacements: dict[str, str], addl_to_remove: set[str], preserve_color: bool, force_all_black: bool) -> bytes | None:
    """
    Replaces placeholders ([[KEY]]) in DOCX content using python-docx.
    """
    logger = logging.getLogger(__name__)
    logging.info("Performing replacements using python-docx...")
    if force_all_black: logger.info("Will force all text runs to black after replacements.")
    elif preserve_color: logger.info("Attempting to preserve original font color.")
    else: logger.info("Forcing color to black for modified runs.")
        
    try:
        doc_stream = io.BytesIO(content_bytes)
        document = docx.Document(doc_stream)

        all_replacements = {**defined_replacements}
        for key in addl_to_remove:
             all_replacements[key] = "" 

        if not all_replacements:
            logger.info("No defined replacements or ADDL_ variables to remove.")
            return content_bytes

        # Find default character style if needed for reset
        default_style = None
        try:
            default_style = document.styles['Default Paragraph Font']
            if default_style.type != WD_STYLE_TYPE.CHARACTER:
                logger.warning("'Default Paragraph Font' is not a Character style.")
                default_style = None # Don't use it if it's not a character style
        except KeyError:
            logger.warning("'Default Paragraph Font' style not found.")


        # Helper function
        def process_paragraph_runs(paragraph):
            for run in paragraph.runs:
                original_text = run.text
                text_to_modify = run.text
                replacement_done_in_run = False
                
                original_color_rgb = run.font.color.rgb
                original_theme_color = run.font.color.theme_color

                for key, value in all_replacements.items():
                     # *** USE NEW DELIMITER FORMAT ***
                     placeholder = f"[[{key}]]" 
                     if placeholder in text_to_modify:
                         text_to_modify = text_to_modify.replace(placeholder, value)
                         replacement_done_in_run = True
                
                if replacement_done_in_run:
                     logger.debug(f"  Replacing text in run: '{original_text[:30]}...' -> '{text_to_modify[:30]}...'")
                     run.text = text_to_modify 

                     # Apply color based on flag
                     if preserve_color and not force_all_black:
                          logger.debug(f"  Attempting preserve color: RGB={original_color_rgb}, Theme={original_theme_color}")
                          run.font.color.rgb = None 
                          run.font.color.theme_color = None
                          if original_color_rgb is not None:
                              try: run.font.color.rgb = RGBColor(original_color_rgb[0], original_color_rgb[1], original_color_rgb[2])
                              except Exception as color_ex: logger.warning(f"Could not apply RGB color {original_color_rgb}: {color_ex}")
                          elif original_theme_color is not None:
                              try: run.font.color.theme_color = original_theme_color
                              except Exception as theme_ex: logger.warning(f"Could not copy theme color ({original_theme_color}): {theme_ex}.")

                     elif not force_all_black:
                          # Force color to black AND attempt to reset character style
                          logger.debug(f"  Forcing color to black and resetting style for run.")
                          if default_style:
                               try:
                                   run.style = default_style
                                   logger.debug(f"    Applied style '{default_style.name}' to run.")
                               except Exception as style_ex:
                                    logger.warning(f"    Failed to apply style '{default_style.name}': {style_ex}")
                          else:
                               logger.debug("    Default character style not available, cannot reset run style.")
                               
                          run.font.color.rgb = RGBColor(0, 0, 0)
                          run.font.color.theme_color = None
                     # else: do nothing here, let final loop handle it.

        # Process paragraphs and tables for replacements
        if all_replacements:
            for para in document.paragraphs: process_paragraph_runs(para)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs: process_paragraph_runs(para)
            logging.info("Run-level replacements applied.")
        else:
             logging.info("No placeholder replacements to perform.")


        # --- NEW: Force all runs to black if flag is set ---
        if force_all_black:
            logging.info("Applying --force-all-black...")
            black_color = RGBColor(0, 0, 0)
            run_count = 0
            # Iterate through paragraphs
            for para in document.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = black_color
                    run.font.color.theme_color = None # Ensure theme color is unset
                    if run.style and hasattr(run.style, "font") and hasattr(run.style.font, "color"):
                      try:
                        run.style.font.color.rgb = black_color
                        run.style.font.color.theme_color = None
                        logging.debug(f"Painted it Black '{run.style.name}'.")
                      except Exception as style_ex:
                        logging.warning(f"Could not force style color to black for style '{run.style.name}'")
                    run_count += 1
            # Iterate through tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.color.rgb = black_color
                                run.font.color.theme_color = None
                                if run.style and hasattr(run.style, "font") and hasattr(run.style.font, "color"):
                                  try:
                                    run.style.font.color.rgb = black_color
                                    run.style.font.color.theme_color = None
                                    logging.debug(f"Painted it Black '{run.style.name}'.")
                                  except Exception as style_ex:
                                    logging.warning(f"Could not force style color to black for '{run.style.name}'")

                                run_count += 1
            logging.info(f"Forced color to black for {run_count} runs in paragraphs and tables.")
                                   
        # Save the modified document
        output_stream = io.BytesIO()
        document.save(output_stream)
        output_stream.seek(0)
        logging.info("Run-level replacements applied successfully.")
        return output_stream.getvalue()

    except Exception as e:
        logging.error(f"Error replacing placeholders using python-docx: {e}", exc_info=True)
        return None 


# --- File Upload/Download/Delete Operations ---
# (upload_temp_file remains the same)
async def upload_temp_file(client: GraphServiceClient, drive_id: str, parent_folder_id: str, filename: str, content: bytes) -> str | None:
    logger = logging.getLogger(__name__)
    temp_item_id = None
    encoded_filename = urllib.parse.quote(filename)
    logger.debug(f"Attempting to upload temporary file '{filename}' (encoded: '{encoded_filename}') to parent {parent_folder_id} in drive {drive_id} ({len(content)} bytes)")
    try:
        target_item_specifier = f"{parent_folder_id}:/{encoded_filename}:"
        logger.debug(f"Using target item specifier for upload PUT: '{target_item_specifier}'")
        response = await client.drives.by_drive_id(drive_id).items.by_drive_item_id(target_item_specifier).content.put(content)
        if response and response.id:
             temp_item_id = response.id
             logger.info(f"Successfully uploaded temporary file '{filename}' with ID: {temp_item_id}")
        else:
             logger.error(f"Upload call succeeded but response did not contain item ID for '{filename}'. Response: {response}")
    except ODataError as o_data_error:
        logger.error(f"Error uploading temporary file '{filename}':")
        if o_data_error.error:
            logger.error(f"  Code: {o_data_error.error.code}")
            logger.error(f"  Message: {o_data_error.error.message}")
            if "invalidRequest" in (o_data_error.error.code or "") or "malformed" in (o_data_error.error.message or ""):
                 logger.error(f"  Hint: The path specifier '{target_item_specifier}' might be incorrectly formed or encoded.")
        logger.debug("ODataError details:", exc_info=True)
    except Exception as e:
        logger.exception(f"An unexpected error occurred uploading temporary file '{filename}': {e}")
    return temp_item_id

# (Workaround Class remains the same)
@dataclass
class ContentFormatQueryParameters(QueryParameters):
    format: str | None = None

# (download_as_pdf remains the same)
async def download_as_pdf(client: GraphServiceClient, drive_id: str, item_id: str, output_file_path: str) -> bool:
    logger = logging.getLogger(__name__)
    logger.debug(f"Requesting PDF conversion for item {item_id} in drive {drive_id} using manual URL workaround...")
    graph_base_url = "https://graph.microsoft.com/v1.0"
    target_url = f"{graph_base_url}/drives/{drive_id}/items/{item_id}/content?format=pdf"
    logger.debug(f"Constructed URL: {target_url}")
    try:
        global auth_credential
        if not auth_credential:
             logger.error("Authentication credential not available for manual request.")
             return False
        logger.debug("Attempting to get token for manual request (synchronously)...")
        loop = asyncio.get_running_loop()
        token = await loop.run_in_executor(
            None, lambda: auth_credential.get_token("https://graph.microsoft.com/.default")
        )
        if not token or not token.token:
             logger.error("Failed to get token for manual request.")
             return False
        logger.debug("Successfully obtained token for manual request.")
        headers = {'Authorization': f'Bearer {token.token}'}
        logger.debug(f"Making GET request to: {target_url}")
        response = await loop.run_in_executor(
            None, lambda: requests.get(target_url, headers=headers, allow_redirects=False)
        )
        logger.debug(f"Initial request completed with status: {response.status_code}")
        if response.status_code == 302:
            download_url = response.headers.get('Location')
            if not download_url:
                 logger.error("PDF conversion request redirected (302) but Location header was missing.")
                 return False
            logger.info(f"Received pre-authenticated download URL: {download_url[:80]}...")
            logger.debug("Making GET request to pre-authenticated download URL...")
            pdf_response = await loop.run_in_executor(
                None, lambda: requests.get(download_url)
            )
            pdf_response.raise_for_status()
            pdf_bytes = pdf_response.content
            logger.info("Successfully downloaded PDF content.")
        elif response.status_code == 200:
             logger.warning("PDF conversion request returned 200 OK directly, expected 302 Redirect. Content might be raw PDF.")
             pdf_bytes = response.content
        else:
             logger.error(f"PDF conversion request failed. Status: {response.status_code}")
             try: error_body = response.json(); logger.error(f"Error Body: {error_body}")
             except requests.exceptions.JSONDecodeError: logger.error(f"Error Body: {response.text}")
             response.raise_for_status()
             return False
        if not pdf_bytes:
            logger.error("PDF download failed to retrieve content bytes.")
            return False
        try:
            with open(output_file_path, 'wb') as f: f.write(pdf_bytes)
            logger.info(f"Successfully saved PDF to {output_file_path} ({len(pdf_bytes)} bytes)")
            if len(pdf_bytes) < 1000: logger.warning(f"Saved PDF file '{output_file_path}' is very small. Please check its contents.")
            return True
        except IOError as e:
            logger.error(f"Error saving PDF file locally '{output_file_path}': {e}", exc_info=True)
            return False
    except requests.exceptions.RequestException as req_ex:
         logger.error(f"HTTP request error during PDF download/conversion: {req_ex}", exc_info=True)
         return False
    except Exception as e:
        logger.exception(f"An unexpected error occurred during manual PDF download for {item_id}: {e}")
        return False

# (delete_drive_item remains the same)
async def delete_drive_item(client: GraphServiceClient, drive_id: str, item_id: str) -> bool:
    logger = logging.getLogger(__name__)
    logger.debug(f"Attempting to delete item {item_id} from drive {drive_id}")
    try:
        await client.drives.by_drive_id(drive_id).items.by_drive_item_id(item_id).delete()
        logger.info(f"Successfully deleted temporary item {item_id}.")
        return True
    except ODataError as o_data_error:
        logger.error(f"Error deleting item {item_id}:")
        if o_data_error.error:
            logger.error(f"  Code: {o_data_error.error.code}")
            logger.error(f"  Message: {o_data_error.error.message}")
        logger.debug("ODataError details:", exc_info=True)
        return False
    except Exception as e:
        logger.exception(f"An unexpected error occurred deleting item {item_id}: {e}")
        return False


# --- Main Workflow ---
async def main(input_onedrive_path: str, output_local_path: str, definitions: dict[str, str], preserve_color: bool, force_all_black: bool):
    logger = logging.getLogger(__name__) # Get logger for main scope
    client = await get_authenticated_client()
    if not client:
        logger.critical("Exiting due to authentication failure.")
        return

    logger.info(f"Processing template file: {input_onedrive_path}")
    original_item_id, parent_folder_id, drive_id = await get_drive_item_details(client, input_onedrive_path)
    if not original_item_id or not parent_folder_id or not drive_id:
        logger.critical("Failed to get template details (ID, Parent ID, or Drive ID). Exiting.")
        return

    logger.info(f"Downloading template content (Item ID: {original_item_id})...")
    original_content = await get_file_content(client, drive_id, original_item_id)
    if original_content is None:
        logger.critical("Failed to retrieve template content. Exiting.")
        return
    if len(original_content) == 0:
         logger.warning("Template content downloaded is empty.")

    logger.info("Scanning template for placeholders {{...}}...")
    found_placeholders = find_placeholders_in_docx(original_content)
    addl_to_remove = set()
    reportable_undefined = set()
    if found_placeholders:
         logger.info(f"Found placeholders: {', '.join(sorted(list(found_placeholders)))}")
         defined_keys = set(definitions.keys())
         for ph in found_placeholders:
              if ph not in defined_keys:
                   if ph.startswith("ADDL_"): addl_to_remove.add(ph)
                   else: reportable_undefined.add(ph)
         if reportable_undefined:
             logger.warning("--- Undefined Placeholders Found ---")
             logger.warning("The following placeholders were found but not defined via -D:")
             for ph in sorted(list(reportable_undefined)): logger.warning(f"  - [[{ph}]]")
             logger.warning("These placeholders will remain unchanged in the output PDF.")
         else: logger.info("All found non-ADDL_ placeholders have definitions provided via -D.")
         if addl_to_remove: logger.info(f"The following undefined ADDL_ placeholders will be removed: {', '.join(sorted(list(addl_to_remove)))}")
    else: logger.info("No placeholders like [[...]] found in the template (or parsing failed).")

    updated_content_or_none = replace_placeholders(original_content, definitions, addl_to_remove, preserve_color, force_all_black)
    if updated_content_or_none is None:
         logger.critical("Failed to replace placeholders in the document. Exiting.")
         return
    updated_content = updated_content_or_none

    original_path = Path(input_onedrive_path)
    temp_filename = f"{original_path.stem}_temp_{uuid.uuid4().hex}{original_path.suffix}"
    logger.info(f"Generated temporary filename: {temp_filename}")

    logger.info(f"Uploading temporary file with replaced content to Parent ID: {parent_folder_id}...")
    temp_item_id = await upload_temp_file(client, drive_id, parent_folder_id, temp_filename, updated_content)

    if not temp_item_id:
        logger.critical("Failed to upload temporary file. Exiting.")
        return

    pdf_download_successful = False
    local_save_successful = False
    try:
        logger.info("Waiting briefly before requesting conversion...")
        await asyncio.sleep(5)
        logger.info(f"Starting PDF download process for temporary item {temp_item_id}...")
        local_save_successful = await download_as_pdf(client, drive_id, temp_item_id, output_local_path)
        pdf_download_successful = local_save_successful
    finally:
        if temp_item_id:
            if local_save_successful:
                 logger.info(f"Attempting to delete temporary file {temp_item_id}...")
                 delete_success = await delete_drive_item(client, drive_id, temp_item_id)
                 if not delete_success: logger.warning(f"Failed to delete temporary file {temp_item_id}. Manual cleanup may be required.")
            else: logger.warning(f"Skipping deletion of temporary file {temp_item_id} because PDF generation/saving failed.")

    if local_save_successful: logger.info("Cover letter generation complete.")
    else: logger.error("Cover letter generation failed.")

# --- Entry Point ---
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate a customized cover letter from a OneDrive template and save as local PDF.",
        formatter_class=argparse.RawTextHelpFormatter
        )
    # (Arguments remain the same)
    parser.add_argument(
        "-i", "--input",
        required=True,
        help="OneDrive path to the input template (.docx) file (e.g., 'Documents/CoverLetterTemplate.docx')"
        )
    parser.add_argument(
        "-o", "--output",
        required=True,
        help="Local file path to save the output PDF (e.g., 'MyCoverLetter.pdf')"
        )
    parser.add_argument(
        "-D", "--define",
        metavar="KEY=VALUE",
        nargs='+',
        action=DefineAction,
        dest='definitions',
        default={},
        help="Define placeholder replacements. Use the format KEY=VALUE.\n"
             "The script will replace occurrences of [[KEY]] in the template with VALUE.\n"
             "Multiple -D arguments can be provided, or multiple KEY=VALUE pairs after one -D.\n"
             "Example: -D COMPANY=\"Example Inc.\" -D ATTN_NAME=\"Ms. Smith\""
        )
    parser.add_argument(
        "-v", "--verbose",
        action="store_true", # Simple flag for debug mode
        help="Enable verbose (DEBUG level) logging."
        )
    parser.add_argument(
        "--preserve-color",
        action="store_true",
        dest="preserve_color",
        default=False,
        help="Attempt to preserve original font color during replacement instead of forcing black."
        )
    parser.add_argument(
        "--force-all-black",
        action="store_true",
        dest="force_all_black",
        default=False,
        help="Force ALL text runs in the document (paragraphs and tables) to black after replacements."
        )
    args = parser.parse_args()

    # --- Configure Logging Level Based on Args ---
    logger = logging.getLogger() # Get root logger
    az_identity_logger = logging.getLogger('azure.identity') # Get azure-identity logger

    if args.verbose:
        logger.setLevel(logging.DEBUG) # Set root to DEBUG
        az_identity_logger.setLevel(logging.INFO)  # Show azure-identity INFO logs in verbose mode
        logging.info("Verbose logging enabled (including azure.identity INFO).")
    else:
        logger.setLevel(logging.INFO) # Set root to INFO
        az_identity_logger.setLevel(logging.WARNING) # Keep azure-identity quiet by default
    # --- End Logging Level Config ---


    # --- Output Directory Handling ---
    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            logger.info(f"Created output directory: {output_dir}") # Use logger
        except OSError as e:
            logger.critical(f"Error: Could not create output directory '{output_dir}': {e}")
            sys.exit(1)

    if not args.output.lower().endswith(".pdf"):
         logger.warning(f"Output file '{args.output}' does not end with .pdf") # Use logger

    # --- Run Main Async Function ---
    try:
        logger.info("Starting main execution.")
        asyncio.run(main(args.input, args.output, args.definitions, args.preserve_color, args.force_all_black))
        logger.info("Main execution finished.")
    except KeyboardInterrupt:
         logger.warning("\nOperation cancelled by user.")
         sys.exit(1)
    except Exception as e:
        logger.critical(f"An unexpected critical error occurred during execution.", exc_info=True)
        sys.exit(1)
